# Dashboard Last Version

import streamlit as st
import pandas as pd
import openai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os
import re
import tempfile
import zipfile
from io import BytesIO
from docx.shared import Pt

# وابستگی‌های اضافه شده برای کراولینگ
import requests
from bs4 import BeautifulSoup
from googlesearch import search
from urllib.parse import urlparse

# ---------------------------------------------
# تابع احراز هویت ساده
# ---------------------------------------------

def login():
    # بررسی می‌کنیم آیا کاربر قبلاً وارد شده یا خیر
    if "logged_in" not in st.session_state:
        st.session_state["logged_in"] = False
    if not st.session_state["logged_in"]:
        st.sidebar.subheader("ورود به داشبورد")
        username = st.sidebar.text_input("نام کاربری")
        password = st.sidebar.text_input("رمز عبور", type="password")
        if st.sidebar.button("ورود"):
            # تغییر دهید: نام کاربری و رمز عبور دلخواه خود را قرار دهید
            if username == "user" and password == "pass":
                st.session_state["logged_in"] = True
                st.sidebar.success("ورود موفقیت‌آمیز!")
            else:
                st.sidebar.error("نام کاربری یا رمز عبور اشتباه است.")
        return st.session_state["logged_in"]
    else:
        return True

# ---------------------------------------------
# توابع کمکی استخراج اطلاعات از اکسل
# ---------------------------------------------

def extract_keywords_from_row(row):
    """ استخراج کلمات کلیدی اصلی و فرعی از هر ردیف اکسل """
    separator_pattern = r"[،;|؛]+"
    keywords_main = re.split(separator_pattern, str(row.get("کلمه کلیدی اصلی", "") or ""))
    keywords_main = [kw.strip() for kw in keywords_main if kw.strip()]
    keywords_sub = re.split(separator_pattern, str(row.get("کلمات کلیدی فرعی", "") or ""))
    keywords_sub = [kw.strip() for kw in keywords_sub if kw.strip()]
    return keywords_main, keywords_sub

def extract_prompt_data(row):
    """ استخراج اطلاعات مورد نیاز برای تولید پرامپت از هر ردیف اکسل """
    original_word_count = row.get("تعداد کلمه", 2000)
    double_word_count = original_word_count * 1  # افزایش 1.3 (در صورت نیاز تغییر دهید)
    return {
        "موضوع": row.get("موضوع", "موضوع مشخص نشده"),
        "نوع مقاله": row.get("نوع مقاله", "نوع مقاله مشخص نشده"),
        "تعداد کلمه": double_word_count,
        "کلمات کلیدی": row.get("کلمات کلیدی", None),
        "لینک‌ها": [row.get("لینک1"), row.get("لینک2")],         # لینک‌ها از اکسل
        "انکرتکست": [row.get("انکرتکست1"), row.get("انکرتکست2")], # انکرتکست‌های مربوط به لینک‌ها
        "عنوان اصلی": row.get("عنوان اصلی(H1)", "عنوان خلاقانه مقاله"),
        "عناوین H2": row.get("عناوین H2", None),
        "عناوین H3": row.get("عناوین H3", None)
    }

# ---------------------------------------------
# توابع کمکی مربوط به حذف URLها از مطالب کراول شده
# ---------------------------------------------

def remove_urls(text):
    """
    تمام URLهای موجود در متن را حذف می‌کند.
    """
    return re.sub(r'https?://\S+', '', text)

# ---------------------------------------------
# توابع مربوط به تولید پرامپت و ارسال به OpenAI
# ---------------------------------------------

def generate_dynamic_prompt(data, keywords_main, keywords_sub):
    """ تولید پرامپت پویا بر اساس اطلاعات اکسل """
    if data["نوع مقاله"] == "سفرنامه":
        prompt = f"لطفاً یک سفرنامه فارسی برای موضوع زیر بنویس:\n"
        prompt += f"- موضوع: {data['موضوع']}\n"
        prompt += f"- تعداد کلمات: حدود {data['تعداد کلمه']}\n"

        if keywords_main:
            prompt += f"- کلمات کلیدی اصلی: {', '.join(keywords_main)}\n"

        if keywords_sub:
            prompt += f"- کلمات کلیدی فرعی: {', '.join(keywords_sub)}\n"

        if any(data["لینک‌ها"]):
            prompt += "- لینک‌ها:\n"
            for link in data["لینک‌ها"]:
                if link:
                    prompt += f"  - {link}\n"
        prompt += "\nنکات راهنما:\n"
        prompt += "- متن سفرنامه باید از زبان اول شخص نوشته شود و لحن آن صمیمی و جذاب باشد.\n"
        prompt += "- تجربیات شخصی، تصمیم‌گیری‌ها و اتفاقات روزانه سفر را بیان کن.\n"
        prompt += "- در متن سفرنامه، از هر کلمه کلیدی اصلی حداقل 3 بار و از هر کلمه کلیدی فرعی حداقل 2 بار(در صورت وجود) به‌طور طبیعی و مرتبط با موضوع در متن استفاده کن.\n"
        prompt += "- داستان‌هایی مانند رزرو بلیط، انتخاب هتل، برخورد با مردم محلی و لحظات خاص سفر را اضافه کن.\n"
        prompt += "- برای هر عنوان H1 حداقل دو پاراگراف طولانی بنویس و جذاب باشد.\n"
        prompt += "- برای هر عنوان H2 و H3 حداقل یک پاراگراف خیلی طولانی در حداقل 5 سطر آماده کند و جذاب باشد.\n"
        prompt += "- عناوین H2 و H3 باید منطقی و مرتبط باشند و حتما بین آنها متن مرتبط باشد . برای هر H2، حداقل 2 تا 4 عنوان H3 نوشته شود.\n"
        prompt += "- حتما در توضیحات هر عنوان از اسامی افراد یا اماکن یا خوراکی‌های مرتبط، تاریخچه، سبک معماری، سال‌هایی که اتفاقات مهمی در مکانی یا شهری یا آثار تاریخی افتاده و به طور کلی جزئیات کاربردی مرتبط با هر عنوان استفاده کن، البته اگر اطلاعات درستی در آن موارد داری.\n"
        prompt += "- لینک‌ها باید کلیک‌پذیر باشند و در متن به صورت طبیعی قرار گیرند.\n"
        prompt += "- کلمات یا جملات را به هیچ عنوان بولد نکن و از شماره گذاری عناوین خودداری کن.\n"
        prompt += "- اطلاعات دقیق و مستند ارائه کن. اگر اطلاعات دقیق ندارید، از گمانه‌زنی خودداری کن.\n"
        prompt += "- یک نکته مهم که حتما باید رعایت شود: (اگر کلمه کلیدی اصلی یا کلمات کلیدی فرعی عباراتی بودند مانند: 'هتل قشم'، 'هتل رشت'، 'هتل چالوس'، 'هتل اصفهان'، 'هتل کاشان' و مانند این موارد)، ترکیب (هتل'نام شهر')، وجود خارجی ندارند، یعنی هتلی به اسم هتل رشت یا هتل اصفهان یا هر ترکیب (هتل'نام شهر') در قالب هتل، وجود خارجی ندارند و منظور از آنها هتلی با آن نام نیست بلکه این کلمات به خاطر اینکه زیاد سرچ میشوند به عنوان کلمه کلیدی در نظر گرفته شدند. از آنها به‌طور خلاقانه در متن استفاده کن، به‌عنوان مثال: 'برای رزرو ارزانترین هتل چالوس به سایت اسنپ‌تریپ مراجعه کنید.' یا 'اگر به دنبال اقامتی آرام در شمال ایران هستید، پیشنهاد می‌کنیم با جستجوی 'هتل چالوس' یا 'هتل رامسر' در سایت‌های معتبر، گزینه‌های خود را مقایسه کنید.' یا 'بزرگترین هتل قشم در مرکز جزیره واقع شده.' یا 'نزدیکترین هتل کاشان به باغ فین هتل(نام نزدیکترین هتل به باغ فین) است.') این جمله باید هربار به صورت رندوم با استفاده از ساختار یکی از نمونه‌ها نوشته شود و به گونه‌ای باشد که در بافت طبیعی متن قرار گیرد و منطقی باشد.\n" 
        prompt += "- به هیچ عنوان عباراتی با ساختار (هتل'نام شهر') رو به عنوان هتل در نظر نگیر.رعایت این مورد بسیار مهم هست.\n"
        prompt += "- بخش نتیجه‌گیری با عنوانی خلاقانه و متناسب با موضوع باشد و از استفاده از عبارات کلیشه‌ای مانند 'جمع‌بندی' و 'نتیجه‌گیری' در عنوان خودداری شود.\n"
        prompt += "- نوع نگارش: \"Human-like, engaging, detailed, SEO-friendly\"\n"
        prompt += "- عناوین باید با تعداد `#`های مناسب شروع شوند و بعد از `#`ها یک فاصله وجود داشته باشد.\n"

    else:
        # پرامپت عمومی برای سایر مقالات
        prompt = f"لطفاً یک مقاله فارسی برای موضوع زیر بنویس:\n"
        prompt += f"- موضوع: {data['موضوع']}\n"
        prompt += f"- نوع مقاله: {data['نوع مقاله']}\n"
        prompt += f"- تعداد کلمات: حدود {data['تعداد کلمه']}\n"

        if keywords_main:
            prompt += f"- کلمات کلیدی اصلی: {', '.join(keywords_main)}\n"

        if keywords_sub:
            prompt += f"- کلمات کلیدی فرعی: {', '.join(keywords_sub)}\n"

        if any(data["لینک‌ها"]):
            prompt += "- لینک‌ها:\n"
            for link in data["لینک‌ها"]:
                if link:
                    prompt += f"  - {link}\n"

        prompt += "\nنکات راهنما:\n"
        prompt += "- لحن مقاله باید متناسب با نوع مقاله باشد.\n"
        prompt += "- از هر کلمه کلیدی اصلی حداقل 3 بار و از هر کلمه کلیدی فرعی حداقل 2 بار(در صورت وجود) به‌طور طبیعی و مرتبط با موضوع در متن استفاده کن.\n"
        prompt += "- برای هر عنوان H1 حداقل دو پاراگراف طولانی بنویس و جذاب باشد.\n"
        prompt += "- برای هر عنوان H2 و H3 حداقل یک پاراگراف خیلی طولانی در حداقل 5 سطر آماده کند و جذاب باشد.\n"
        prompt += "- هر H2 باید حداقل بین 2 الی 4 H3 داشته باشد.\n"
        prompt += "- لینک‌ها باید کلیک‌پذیر باشند و در متن به صورت طبیعی قرار گیرند.\n"
        prompt += "- کلمات یا جملات را به هیچ عنوان بولد نکن و از شماره گذاری عناوین خودداری کن.\n"
        prompt += "- حتما در توضیحات هر عنوان از اسامی افراد یا اماکن یا خوراکی‌های مرتبط، تاریخچه، سبک معماری، سال‌هایی که اتفاقات مهمی در مکانی یا شهری یا آثار تاریخی افتاده و به طور کلی جزئیات کاربردی مرتبط با هر عنوان استفاده کن، البته اگر اطلاعات درستی در آن موارد داری.\n"
        prompt += "- کلمات یا جملات را به هیچ عنوان بولد نکن و از شماره گذاری عناوین خودداری کن.\n"
        prompt += "- اطلاعات دقیق و مستند ارائه کن. اگر اطلاعات دقیق ندارید، از گمانه‌زنی خودداری کن.\n"
        prompt += "- یک نکته مهم که حتما باید رعایت شود: (اگر کلمه کلیدی اصلی یا کلمات کلیدی فرعی عباراتی بودند مانند: 'هتل قشم'، 'هتل رشت'، 'هتل چالوس'، 'هتل اصفهان'، 'هتل کاشان' و مانند این موارد)، ترکیب (هتل'نام شهر')، وجود خارجی ندارند، یعنی هتلی به اسم هتل رشت یا هتل اصفهان یا هر ترکیب (هتل'نام شهر') در قالب هتل، وجود خارجی ندارند و منظور از آنها هتلی با آن نام نیست بلکه این کلمات به خاطر اینکه زیاد سرچ میشوند به عنوان کلمه کلیدی در نظر گرفته شدند. از آنها به‌طور خلاقانه در متن استفاده کن، به‌عنوان مثال: 'برای رزرو ارزانترین هتل چالوس به سایت اسنپ‌تریپ مراجعه کنید.' یا 'اگر به دنبال اقامتی آرام در شمال ایران هستید، پیشنهاد می‌کنیم با جستجوی 'هتل چالوس' یا 'هتل رامسر' در سایت‌های معتبر، گزینه‌های خود را مقایسه کنید.' یا 'بزرگترین هتل قشم در مرکز جزیره واقع شده.' یا 'نزدیکترین هتل کاشان به باغ فین هتل(نام نزدیکترین هتل به باغ فین) است.') این جمله باید هربار به صورت رندوم با استفاده از ساختار یکی از نمونه‌ها نوشته شود و به گونه‌ای باشد که در بافت طبیعی متن قرار گیرد و منطقی باشد.\n" 
        prompt += "- به هیچ عنوان عباراتی با ساختار (هتل'نام شهر') رو به عنوان هتل در نظر نگیر.رعایت این مورد بسیار مهم هست.\n"
        prompt += "- بخش نتیجه‌گیری با عنوانی خلاقانه و متناسب با موضوع باشد و از استفاده از عبارات کلیشه‌ای مانند 'جمع‌بندی' و 'نتیجه‌گیری' در عنوان خودداری شود.\n"
        prompt += "- تمام اصول مرتبط با SEO را رعایت کن.\n"
        prompt += "- نوع نگارش: \"Human-like, engaging, detailed, SEO-friendly\"\n"
        prompt += "- حتما در بین عناوین H1 و H2 و همچنین H2 و H3 یک متن مرتبط که حداقل یک پاراگراف طولانی و بلند باشه، قرار بده..\n"
        prompt += "- عناوین باید با تعداد `#`های مناسب شروع شوند و بعد از `#`ها یک فاصله وجود داشته باشد.\n"


    return prompt

def generate_article(prompt, api_key):
    """ ارسال پرامپت به OpenAI و دریافت مقاله """
    openai.api_key = api_key
    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",  
        messages=[
                {
                    "role": "system",
                    "content": (
                        "شما یک نویسنده ریزبین و حرفه‌ای فارسی در حوزه گردشگری هستید که مقالات جامع، خلاقانه و با ذکر جزئیات "
                        "(استفاده از نام‌ها، اعداد، تاریخ‌ها و ویژگی‌های مرتبط با موضوع هر مقاله) تولید می‌کنید. "
                        "مقالات شما باید شامل اطلاعات دقیق و کاربردی باشند که مخاطب را جذب کرده و او را در مسیر تجربه‌های سفر همراهی کنند. "
                        "جزئیات مهم و جالب مانند 'تاریخچه'، 'سبک معماری'، 'مکان‌های مرتبط'، 'آدرس اماکن و جاذبه‌ها' و هر گونه اطلاعات تکمیلی که مقاله را غنی‌تر کند، الزامی است."
                        "بخش نتیجه‌گیری با عنوانی خلاقانه و متناسب با موضوع باشد و از استفاده از عبارات کلیشه‌ای مانند 'جمع‌بندی' و 'نتیجه‌گیری' در عنوان خودداری شود."
                        "برای هر عنوان H1 حداقل دو پاراگراف طولانی بنویس و جذاب باشد."
                        "برای هر عنوان H2 و H3 حداقل یک پاراگراف خیلی طولانی در حداقل 5 سطر آماده کند و جذاب باشد."
                        "هر H2 باید حداقل بین 2 الی 4 H3 داشته باشد."
                        "از 'تکرار کلمات و جملات و مطالب خودداری کن' و متن رو سرشار از اطلاعات \"Unique\" کن."
                        
                        "\n\n"
                        "یک نکته بسیار مهم: اگر ترکیب 'هتل{نام شهر}' (مانند: هتل قشم، هتل رشت، هتل چالوس، هتل اصفهان و هتل کاشان) در کلمات کلیدی اصلی یا فرعی بود، باید حتماً در متن استفاده شود. "
                        "اما دقت کنید: این ترکیب‌ها به‌عنوان نام هتل واقعی وجود ندارند و نباید در متن به‌عنوان نام یک هتل خاص استفاده شوند. "
                        "این کلمات فقط به‌عنوان کلمات کلیدی پرتکرار در نظر گرفته شده‌اند. "
                        "به جای استفاده از این ترکیب‌ها به‌طور خلاقانه، طبیعی و توصیفی در متن استفاده کنید."
                        "\n\n"
                        "مثال‌ها برای استفاده صحیح که هربار به صورت رندوم باید از سبک یکی از مثالها استفاده کنی:\n"
                        "1. 'برای رزرو ارزان‌ترین هتل چالوس به سایت اسنپ‌تریپ مراجعه کنید.'\n"
                        "2. 'اگر به دنبال اقامتی آرام در شمال ایران هستید، پیشنهاد می‌کنیم با جستجوی 'هتل چالوس' یا 'هتل رامسر' در سایت‌های معتبر، گزینه‌های خود را مقایسه کنید.'\n"
                        "3. 'بزرگترین 'هتل قشم' در مرکز جزیره واقع شده.'\n"
                        "4. 'نزدیک‌ترین 'هتل کاشان' به باغ فین هتل(نام نزدیک‌ترین هتل به باغ فین) است.'\n"
                        "5. 'برای تجربه‌ای به‌یادماندنی در سفر به اصفهان، می‌توانید نزدیکترین هتل اصفهان به میدان نقش جهان را جستجو کنید.'\n"
                        "6. '(قدیمی‌ترین یا جدیدترین یا لوکس‌ترین و مانند این صفات توصیفی) 'هتل رشت' در نزدیکی بلوار گلسار قرار دارد.'\n"
                        "7. 'اگر به دنبال بهترین اقامتگاه‌ها در قشم هستید، می‌توانید در صفحه \"هتل قشم\" جستجو کنید.'\n"
                        "8. 'بهترین هتل چالوس را میتوانید در پلتفرم‌های معتبر رزرو هتل، جستجو کنید.'\n"
                        "مثال‌ها برای انواع مقاله که می‌تونی متناسب با نوع مقاله از سبک نگارش استفاده شده در نمونه‌ها استفاده کنی:\n"
                        "1. 'معرفی هتل: \n"
                        "   هتل اسپیناس پالاس یکی از لوکس‌ترین هتل‌های تهران و کشور است که در منطقه سعادت‌آباد تهران قرار دارد. این هتل در سال ۱۳۹۴ افتتاح شد و در دل خود ۲۰ طبقه با ۴۰۰ اتاق و سوییت مجلل را جای داده است. پس از ورود به لابی هتل، شکوه و معماری کلاسیک آن شما را مجذوب خود خواهد کرد. این هتل به عنوان یکی از بهترین هتل‌های تهران با خدماتی از قبیل چندین رستوران و کافه، یک مرکز بیزنس برای مسافرت‌های کاری، سالن بدن‌سازی و استخر سرپوشیده شناخته می‌شود. در کنار این امکانات رفاهی، هتل اسپیناس پالاس از نظر دسترسی به مراکز تجاری و تفریحی در تهران نیز موقعیت مناسبی دارد. فاصله هتل از فرودگاه مهرآباد ۱۸ کیلومتر و از فرودگاه امام خمینی ۶۴ کیلومتر است. خدمات رفاهی این هتل شامل روم سرویس ۲۴ ساعته، اینترنت رایگان، پارکینگ و خدمات شست‌وشو است. برای مسافرانی که به دنبال اقامتی لوکس با تمامی امکانات رفاهی هستند، این هتل انتخابی بی‌نظیر خواهد بود. \n"
                        "   آدرس: سعادت‌آباد، میدان بهرود، خیابان عابدی، کوچه ۳۳.\n"
                        "2. 'معرفی جاذبه‌ها: \n"
                        "   میدان نقش جهان؛ جایی که همیشه دلتنگش می‌شوید. این میدان با ویژگی‌های منحصربه‌فرد خود، یکی از مهم‌ترین جاذبه‌های گردشگری اصفهان است. در گذشته این میدان با نام «میدان شاه» شناخته می‌شد که پس از انقلاب نام آن به «میدان امام» تغییر یافت. این میدان به ابعاد ۵۶۰ متر طول و ۱۶۰ متر عرض، دومین میدان بزرگ جهان است که در فهرست آثار جهانی یونسکو نیز قرار دارد. نقش جهان از جاهای دیدنی اصفهان است که قدمت آن به دوران صفویه بازمی‌گردد. در گذشته این میدان محل برگزاری جشن‌ها و بازی‌های چوگان بود و از سال ۱۳۱۳ در فهرست آثار ملی ایران قرار گرفت. علاوه بر این، این میدان شامل مجموعه‌ای از بناهای تاریخی همچون مسجد امام، کاخ عالی قاپو، مسجد شیخ لطف‌الله و سردر قیصریه است که هرکدام شاهکاری از معماری ایرانی را به نمایش می‌گذارند. دسترسی به میدان نقش جهان به راحتی از طریق تاکسی، اتوبوس، مترو و حتی پیاده امکان‌پذیر است، اما اگر با خودروی شخصی به آنجا می‌روید، باید برای پارک خودرو از خیابان‌های اطراف استفاده کنید.\n"
                        "   در صورتی که از مترو استفاده می‌کنید، باید در ایستگاه میدان امام حسین پیاده شوید و چند دقیقه پیاده‌روی کنید تا به میدان برسید. اگر اتوبوس می‌خواهید، خطوط مختلفی از جمله میدان امام حسین (ع) – باغ غدیر و میدان امام حسین (ع) – هشت بهشت شما را به ایستگاه امام می‌رسانند. \n"
                        "3. 'معرفی جاذبه‌های طبیعی: \n"
                        "   دره ستارگان در قشم، یکی از جذاب‌ترین و بی‌نظیرترین جاذبه‌های طبیعی ایران است. این دره مجموعه‌ای از کوه‌ها، دره‌ها و صخره‌های شگفت‌انگیز است که طی سال‌ها توسط عوامل طبیعی مانند باد، باران و فرسایش به اشکال مختلفی تبدیل شده‌اند. این شکل‌های خاص گاهی به موجودات عجیب و غریب و حیوانات شبیه می‌شوند، و از همین‌رو مردم بومی منطقه معتقدند که این اشکال بر اثر برخورد یک ستاره با زمین ایجاد شده‌اند. به همین دلیل نام این دره، دره ستارگان است. این مکان یک انتخاب عالی برای علاقه‌مندان به طبیعت و عکاسی است، زیرا هر گوشه‌ای از دره درختان و سنگ‌ها به شکلی منحصر به فرد شکل گرفته است. بازدید از این دره تجربه‌ای فراموش‌نشدنی را برای گردشگران به ارمغان می‌آورد.\n"
                        "   اگر به قشم سفر کردید، بازدید از این جاذبه طبیعی را از دست ندهید. این مکان برای کسانی که به دنبال تجربه‌ای متفاوت از طبیعت ایران هستند، یک مقصد عالی به شمار می‌آید.\n"
                        "سایر نکات مهم:\n"
                        "1. از هر کلمه کلیدی اصلی حداقل 3 بار و از هر کلمه کلیدی فرعی حداقل 2 بار(در صورت وجود) به‌طور طبیعی در متن استفاده کنید.\n"
                        "2. متن باید فاقد گمانه‌زنی باشد و تنها بر اطلاعات دقیق و مستند تمرکز کند.\n"
                        "3. تمامی اصول SEO باید رعایت شود، از جمله توزیع مناسب کلمات کلیدی.\n"
                        "4. از کلمات یا جملات بولدشده استفاده نکنید.\n"
                        "5. حتما در بین عناوین H1 و H2 و همچنین H2 و H3 یک متن مرتبط که حداقل یک پاراگراف طولانی و بلند باشه، قرار بده.\n"
                        "6. در متن از اطلاعات جذاب و خلاقانه استفاده کنید که خواننده را به ادامه مطالعه ترغیب کند."
                        "7. از تکرار زیاد یک کلمه یا عبارت (مانند گردشگر یا گردشگری یا امثال این کلمات) خودداری کن و از مترادف‌های آنها استفاده کن."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
        temperature=0.2,
        max_tokens=4000
    )
    return response.choices[0].message.content

def calculate_word_count(text):
    """ محاسبه تعداد کلمات با در نظر گرفتن نیم‌فاصله """
    words = re.findall(r'\b\w+\b', text.replace('\u200c', ''))
    return len(words)

def reduce_word_count_naturally(article_content, desired_word_count, api_key):
    """ کاهش تعداد کلمات در مقاله به صورت طبیعی با استفاده از مدل """
    revise_prompt = (
        f"مقاله زیر باید به تعداد کلمات حدود {desired_word_count} کاهش یابد.\n\n"
        f"--- متن مقاله فعلی ---\n"
        f"{article_content}\n"
        f"-----------------------\n"
        f"لطفاً متن مقاله را به صورت طبیعی و بدون از دست دادن اطلاعات مهم کاهش دهید تا تعداد کلمات آن به حدود {desired_word_count} برسد."
    )
    try:
        reduced_content = generate_article(revise_prompt, api_key)
        return reduced_content
    except Exception as e:
        st.error(f"خطا در کاهش تعداد کلمات مقاله: {e}")
        return article_content

def add_hyperlink(paragraph, text, url):
    """ ایجاد لینک کلیک‌پذیر در پاراگراف """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_text = OxmlElement('w:t')
    new_text.text = text
    new_run.append(new_text)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

def save_to_word(content, filename, links):
    """ ذخیره متن مقاله در فایل Word """
    doc = Document()
    doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.styles['Normal']._element.get_or_add_pPr().set(qn('w:bidi'), '1')
    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        paragraph = None
        if stripped.startswith('#'):
            match = re.match(r'^(#{1,4})\s*(.*)', stripped)
            if match:
                header_marks = match.group(1)
                header_text = match.group(2)
                header_level = len(header_marks)
                if 1 <= header_level <= 4:
                    paragraph = doc.add_heading(header_text, level=header_level)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph._element.get_or_add_pPr().set(qn('w:bidi'), '1')
                else:
                    paragraph = doc.add_paragraph(stripped)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph._element.get_or_add_pPr().set(qn('w:bidi'), '1')
        else:
            paragraph = doc.add_paragraph()
            if any(link_text in stripped for link_text in links.keys()):
                for link_text, url in links.items():
                    if link_text in stripped:
                        parts = stripped.split(link_text)
                        if parts[0]:
                            paragraph.add_run(parts[0])
                        add_hyperlink(paragraph, link_text, url)
                        if len(parts) > 1 and parts[1]:
                            paragraph.add_run(parts[1])
                        break
            else:
                paragraph.add_run(stripped)
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph._element.get_or_add_pPr().set(qn('w:bidi'), '1')
        if paragraph:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = 1.5
    doc.save(filename)

def check_keywords_in_text(keywords, text):
    """ بررسی حضور کلمات کلیدی در متن """
    missing_keywords = [kw.strip() for kw in keywords if kw.strip() not in text]
    return missing_keywords

def validate_article_structure(content, data):
    """ بررسی ساختار مقاله برای اطمینان از وجود متن بین H2 و H3 """
    lines = content.split('\n')
    previous_heading = None
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('## '):
            previous_heading = 'H2'
        elif stripped.startswith('### '):
            if previous_heading not in ['H2', 'H3', 'Text']:
                return False
            previous_heading = 'H3'
        elif stripped and not stripped.startswith('#'):
            previous_heading = 'Text'
    return True

# ---------------------------------------------
# توابع مربوط به کراولینگ اطلاعات (با تشخیص انعطاف‌پذیر دامنه‌ها)
# ---------------------------------------------

def crawl_google_links(topic, allowed_roots, num_links=5):
    """
    جستجو در گوگل بر اساس موضوع و بازگرداندن لینک‌هایی که از نظر دامنه (بدون توجه به http/https و www)
    با یکی از دامنه‌های مجاز تطبیق دارند.
    
    پارامترها:
      - topic: موضوع یا عبارت جستجو
      - allowed_roots: لیستی از دامنه‌های مجاز (مثلاً "https://www.kojaro.com/")
      - num_links: تعداد لینک‌های مورد نظر (پیش‌فرض ۵)
      
    خروجی:
      - لیستی از لینک‌های یکتا.
    """
    links = []
    try:
        query = topic
        count = 0
        # استفاده از تابع search بدون پارامترهای stop و pause
        for url in search(query):
            count += 1
            if is_allowed_url(url, allowed_roots):
                if url not in links:
                    links.append(url)
            if len(links) >= num_links or count >= 20:
                break
    except Exception as e:
        st.error(f"خطا در جستجوی گوگل: {e}")
    return links

def get_domain(url):
    """
    دامنه URL را استخراج می‌کند و پیشوند "www." را حذف می‌کند.
    """
    try:
        parsed = urlparse(url)
        domain = parsed.netloc.lower()
        if domain.startswith("www."):
            domain = domain[4:]
        return domain
    except Exception:
        return ""

def is_allowed_url(url, allowed_roots):
    """
    بررسی می‌کند آیا دامنه URL با یکی از دامنه‌های موجود در allowed_roots تطبیق دارد.
    تفاوت‌های پروتکل (http/https) و پیشوند "www." نادیده گرفته می‌شوند.
    """
    domain = get_domain(url)
    for root in allowed_roots:
        allowed_domain = get_domain(root)
        if domain == allowed_domain or domain.endswith("." + allowed_domain):
            return True
    return False

def fetch_page_content(url):
    """
    دریافت محتوای یک صفحه وب و استخراج متن آن (اولویت به تگ <article> در صورت وجود).
    """
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, "html.parser")
            article = soup.find("article")
            if article:
                text = article.get_text(separator="\n")
            else:
                text = soup.get_text(separator="\n")
            return text
        else:
            return ""
    except Exception as e:
        st.error(f"خطا در دریافت محتوای {url}: {e}")
        return ""

def gather_crawled_information(topic):
    """
    جستجو در گوگل برای موضوع داده‌شده، کراول کردن ۵ لینک اول از منابع مجاز و بازگرداندن کل محتوای استخراج‌شده.
    خروجی نهایی در یک بخش (expander) نمایش داده می‌شود.
    قبل از ارسال به پرامپت، تمام URLها از مطالب حذف می‌شوند.
    تنها بخشی (مثلاً 500 کاراکتر) از مطالب کراول شده در داشبورد نمایش داده می‌شود.
    """
    allowed_roots = [
        "https://www.alibaba.ir/mag/",
        "https://www.kojaro.com/",
        "https://www.flytoday.ir/blog/",
        "https://safarmarket.com/blog/",
        "https://www.otaghak.com/blog/"
    ]
    links = crawl_google_links(topic, allowed_roots, num_links=5)
    if not links:
        st.warning("هیچ نتیجه‌ای از منابع مجاز پیدا نشد.")
        return ""
    info_list = []
    for link in links:
        content = fetch_page_content(link)
        if content:
            info_list.append(f"منبع: {link}\n{content}")
    crawled_info = "\n\n".join(info_list)
    # حذف URLها از مطالب کراول شده
    crawled_info_clean = remove_urls(crawled_info)
    # نمایش فقط بخشی (500 کاراکتر) از مطالب کراول شده در داشبورد
    display_text = crawled_info_clean[:500] + "..." if len(crawled_info_clean) > 500 else crawled_info_clean
    with st.expander("نمایش بخشی از اطلاعات کراول شده"):
        st.text(display_text)
    return crawled_info_clean

# ---------------------------------------------
# اپلیکیشن Streamlit
# ---------------------------------------------

def main():
    # اضافه کردن بخش ورود (login)
    if "logged_in" not in st.session_state:
        st.session_state["logged_in"] = False
    if not st.session_state["logged_in"]:
        st.sidebar.subheader("ورود به داشبورد")
        username = st.sidebar.text_input("نام کاربری")
        password = st.sidebar.text_input("رمز عبور", type="password")
        if st.sidebar.button("ورود"):
            # تنظیم مقادیر مورد نظر برای ورود؛ به عنوان مثال:
            if username == "user" and password == "pass":
                st.session_state["logged_in"] = True
                st.sidebar.success("ورود موفقیت‌آمیز!")
            else:
                st.sidebar.error("نام کاربری یا رمز عبور اشتباه است.")
        st.stop()
    
    st.set_page_config(page_title="Fahva Article Generator Beta V1.1", layout="wide")
    st.title("Fahva Article Generator alltour V1.1")
    
    st.sidebar.header("تنظیمات")
    api_key = st.sidebar.text_input("کلید API OpenAI", type="password")
    uploaded_file = st.sidebar.file_uploader("آپلود فایل اکسل", type=["xlsx", "xls"])
    model_name = st.sidebar.selectbox("انتخاب مدل OpenAI", ["gpt-4o-mini", "gpt-3.5-turbo"])
    
    if st.sidebar.button("شروع تولید مقالات"):
        if not api_key:
            st.error("لطفاً کلید API OpenAI را وارد کنید.")
            return
        if not uploaded_file:
            st.error("لطفاً فایل اکسل را آپلود کنید.")
            return
        try:
            df = pd.read_excel(uploaded_file, engine='openpyxl')
        except Exception as e:
            st.error(f"خطا در خواندن فایل اکسل: {e}")
            return
        required_columns = ["موضوع", "نوع مقاله", "تعداد کلمه", "کلمه کلیدی اصلی", "کلمات کلیدی فرعی", "لینک1", "انکرتکست1"]
        for col in required_columns:
            if col not in df.columns:
                st.error(f"ستون اجباری '{col}' در فایل اکسل وجود ندارد.")
                return
        with tempfile.TemporaryDirectory() as tmpdirname:
            st.success(f"مسیر موقت برای ذخیره فایل‌ها: {tmpdirname}")
            generated_files = []
            for idx, row in df.iterrows():
                data = extract_prompt_data(row)
                keywords_main, keywords_sub = extract_keywords_from_row(row)
                st.info(f"در حال جستجو و کراول اطلاعات مرتبط با موضوع '{data['موضوع']}' برای ردیف {idx + 1}...")
                crawled_info = gather_crawled_information(data["موضوع"])
                attempt = 0
                max_attempts = 3
                successful = False
                while attempt < max_attempts and not successful:
                    attempt += 1
                    st.write(f"تولید مقاله برای ردیف {idx + 1} - تلاش {attempt}")
                    prompt = generate_dynamic_prompt(data, keywords_main, keywords_sub, crawled_info)
                    try:
                        article_content = generate_article(prompt, api_key)
                        missing_main = check_keywords_in_text(keywords_main, article_content) if keywords_main else []
                        missing_sub = check_keywords_in_text(keywords_sub, article_content) if keywords_sub else []
                        desired_word_count = int(data['تعداد کلمه'])
                        lower_limit = int(desired_word_count * 0.9)
                        upper_limit = int(desired_word_count * 1.1)
                        word_count = calculate_word_count(article_content)
                        if word_count > upper_limit:
                            st.info(f"تعداد کلمات ({word_count}) بیش از حد مجاز ({upper_limit}) است. در حال کاهش تعداد کلمات...")
                            article_content = reduce_word_count_naturally(article_content, desired_word_count, api_key)
                            word_count = calculate_word_count(article_content)
                            st.info(f"تعداد کلمات پس از کاهش: {word_count}")
                        if word_count < lower_limit:
                            extended_word_count = desired_word_count
                            missing_keywords = []
                            if missing_main:
                                missing_keywords.append(f"کلمات کلیدی اصلی: {', '.join(missing_main)}")
                            if missing_sub:
                                missing_keywords.append(f"کلمات کلیدی فرعی: {', '.join(missing_sub)}")
                            revise_prompt = (
                                f"مقاله زیر باید بازنویسی شود تا:\n"
                                f"- تعداد کلمات به حدود {extended_word_count} برسد.\n"
                                f"- کلمات کلیدی زیر که استفاده نشده‌اند، به متن اضافه شوند:\n"
                                f"{'; '.join(missing_keywords)}\n\n"
                                f"--- متن مقاله فعلی ---\n"
                                f"{article_content}\n"
                                f"-----------------------\n"
                                f"لطفاً متن مقاله را گسترش دهید و کلمات کلیدی بالا را به طور طبیعی و مناسب در متن بگنجانید."
                            )
                            try:
                                article_content = generate_article(revise_prompt, api_key)
                                word_count = calculate_word_count(article_content)
                                word_count_valid = lower_limit <= word_count <= upper_limit
                                missing_main = check_keywords_in_text(keywords_main, article_content) if keywords_main else []
                                missing_sub = check_keywords_in_text(keywords_sub, article_content) if keywords_sub else []
                                if not missing_main and not missing_sub and word_count_valid:
                                    successful = True
                                    st.success(f"مقاله ردیف {idx + 1} با موفقیت بازنویسی و تکمیل شد.")
                                else:
                                    st.warning("مقاله بازنویسی شد اما همچنان مشکلاتی دارد: کلمات کلیدی استفاده‌نشده یا تعداد کلمات.")
                            except Exception as e:
                                st.error(f"خطا در بازنویسی مقاله برای ردیف {idx + 1}: {e}")
                                break
                        word_count_valid = lower_limit <= word_count <= upper_limit
                        if missing_main or missing_sub or not word_count_valid:
                            if word_count > upper_limit:
                                if attempt < max_attempts:
                                    st.warning(f"مقاله ردیف {idx + 1} نیاز به کاهش تعداد کلمات دارد.")
                                continue
                            else:
                                extended_word_count = desired_word_count * 2 if not word_count_valid else desired_word_count
                                missing_keywords = []
                                if missing_main:
                                    missing_keywords.append(f"کلمات کلیدی اصلی: {', '.join(missing_main)}")
                                if missing_sub:
                                    missing_keywords.append(f"کلمات کلیدی فرعی: {', '.join(missing_sub)}")
                                revise_prompt = (
                                    f"مقاله زیر باید بازنویسی شود تا:\n"
                                    f"- تعداد کلمات به حدود {extended_word_count} برسد.\n"
                                    f"- کلمات کلیدی زیر که استفاده نشده‌اند، به متن اضافه شوند:\n"
                                    f"{'; '.join(missing_keywords)}\n\n"
                                    f"--- متن مقاله فعلی ---\n"
                                    f"{article_content}\n"
                                    f"-----------------------\n"
                                    f"لطفاً متن مقاله را گسترش دهید و کلمات کلیدی بالا را به طور طبیعی و مناسب در متن بگنجانید."
                                )
                                try:
                                    article_content = generate_article(revise_prompt, api_key)
                                    word_count = calculate_word_count(article_content)
                                    word_count_valid = lower_limit <= word_count <= upper_limit
                                    missing_main = check_keywords_in_text(keywords_main, article_content) if keywords_main else []
                                    missing_sub = check_keywords_in_text(keywords_sub, article_content) if keywords_sub else []
                                    if not missing_main and not missing_sub and word_count_valid:
                                        successful = True
                                        st.success(f"مقاله ردیف {idx + 1} با موفقیت بازنویسی و تکمیل شد.")
                                    else:
                                        st.warning("مقاله بازنویسی شد اما همچنان مشکلاتی دارد: کلمات کلیدی استفاده‌نشده یا تعداد کلمات.")
                                except Exception as e:
                                    st.error(f"خطا در بازنویسی مقاله برای ردیف {idx + 1}: {e}")
                                    break
                        if not missing_main and not missing_sub and word_count_valid:
                            successful = True
                            st.success(f"مقاله ردیف {idx + 1} با موفقیت تولید شد.")
                            st.info(f"تعداد کلمات نهایی مقاله: {word_count}")
                            safe_title = re.sub(r'[\\/*?:"<>|]', "", data['عنوان اصلی']).strip().replace(' ', '_')
                            links = {f"لینک{i+1}": link for i, link in enumerate(data["لینک‌ها"]) if link}
                            output_filename = os.path.join(tmpdirname, f"article_{safe_title}_{idx + 1}.docx")
                            save_to_word(article_content, output_filename, links)
                            generated_files.append((f"مقاله {idx + 1}: {data['عنوان اصلی']}", output_filename))
                            st.write(f"مقاله ذخیره شد: {output_filename}")
                        else:
                            if word_count > upper_limit:
                                st.warning(f"مقاله ردیف {idx + 1} نیاز به کاهش تعداد کلمات دارد.")
                            else:
                                st.warning(f"مقاله ردیف {idx + 1} نیاز به بازنویسی دارد.")
                                if missing_main:
                                    st.warning(f"کلمات کلیدی اصلی استفاده‌نشده: {', '.join(missing_main)}")
                                if missing_sub:
                                    st.warning(f"کلمات کلیدی فرعی استفاده‌نشده: {', '.join(missing_sub)}")
                                if not word_count_valid:
                                    st.warning(f"تعداد کلمات مقاله ({word_count}) در بازه مجاز نیست.")
                    except Exception as e:
                        st.error(f"خطا در تولید مقاله برای ردیف {idx + 1}: {e}")
                if not successful:
                    st.error(f"مقاله ردیف {idx + 1} پس از {max_attempts} تلاش تولید نشد.")
            if generated_files:
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zipf:
                    for file_name, file_path in generated_files:
                        zipf.write(file_path, os.path.basename(file_path))
                zip_buffer.seek(0)
                st.download_button(
                    label="دانلود همه مقالات به صورت ZIP",
                    data=zip_buffer,
                    file_name="generated_articles.zip",
                    mime="application/zip"
                )
            else:
                st.warning("هیچ مقاله‌ای تولید نشد.")

if __name__ == "__main__":
    main()
